"""Render a wrap report as a Word document (#184).

Drew: *"We need to make certain documents human readable, so having a wrap
report as an md doc isn't correct."*

WHY .docx AND NOT A GOOGLE DOC. A Google Doc needs Drive write access, which
lives in mc-2 behind per-user OAuth — and `GoogleDriveClient()` is on record
as HANGING when it has no token. That is a bad failure for a close-out ritual
someone runs once per project. A .docx is just a file: it lands in Dropbox
`03 Assets/06 Spine/` through `push_to_dropbox`, which already works, and
Dropbox previews it in the browser and on mobile without a download. It also
attaches to an email, which a Drive link does not do for an external client.

WHAT THIS IS NOT. It is not a markdown-to-Word converter. It takes the
STRUCTURED bundle plus the model's authored prose and lays out a document
designed to be read by a person who was not on the project — headings, real
tables for the facts, and the human-entry fields rendered as visible blanks
rather than omitted. Feeding it raw markdown would reproduce the exact
problem it exists to fix.

The markdown stays the source of record in the tenant tree (greppable,
diffable, spine-attachable). The .docx is the READABLE projection of it.
Both, not either.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

# Human-entry placeholder. Deliberately conspicuous: the whole point of the
# social-builder split is that an unfilled field should look unfilled.
BLANK = "_____________"


@dataclass
class WrapSection:
    """One authored section of the report.

    `body` is the model's prose. `table` is optional (header row first).
    A section may carry either, both, or neither (a heading-only spacer).
    """

    heading: str
    body: str = ""
    table: list[list[str]] = field(default_factory=list)
    # Rendered as a visible blank line the reader is meant to fill in.
    blanks: list[str] = field(default_factory=list)
    level: int = 1


def _add_table(doc: Any, rows: list[list[str]]) -> None:
    """Write a header-first table in a built-in style.

    `Light Grid Accent 1` is a stock Word style, so the file opens correctly
    on a machine that has never seen our template — the failure mode of
    styling against a custom template we don't ship.
    """
    if not rows:
        return
    table = doc.add_table(rows=1, cols=len(rows[0]))
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:  # pragma: no cover — style set varies by Word build
        pass
    for i, label in enumerate(rows[0]):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(str(label))
        run.bold = True
    for row in rows[1:]:
        cells = table.add_row().cells
        for i, value in enumerate(row[: len(rows[0])]):
            cells[i].text = str(value)


def build_wrap_docx(
    *,
    title: str,
    subtitle: str,
    sections: list[WrapSection],
    out_path: Path,
) -> Path:
    """Write the report and return its path.

    Import is function-local so the whole CLI doesn't pay python-docx's
    import cost on every invocation — only `cp wrap --docx` does.
    """
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading(title, level=0)
    if subtitle:
        para = doc.add_paragraph()
        run = para.add_run(subtitle)
        run.italic = True
        run.font.size = Pt(11)

    for section in sections:
        doc.add_heading(section.heading, level=section.level)
        if section.body:
            for block in section.body.split("\n\n"):
                text = block.strip()
                if not text:
                    continue
                # A leading "- " block becomes real bullets rather than a
                # paragraph that merely starts with a hyphen.
                lines = [ln.strip() for ln in text.split("\n") if ln.strip()]
                if all(ln.startswith(("- ", "* ")) for ln in lines):
                    for ln in lines:
                        doc.add_paragraph(ln[2:].strip(), style="List Bullet")
                else:
                    doc.add_paragraph(text)
        if section.table:
            _add_table(doc, section.table)
        for label in section.blanks:
            para = doc.add_paragraph()
            run = para.add_run(f"{label}: ")
            run.bold = True
            para.add_run(BLANK)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def facts_table(payload: dict) -> list[list[str]]:
    """The at-a-glance facts block, from a `cp wrap --bundle` payload.

    Every value is either a real measured number or an explicit marker. A
    missing number renders as "not recorded" — never as a plausible-looking
    zero, which is how a $0 margin becomes a talking point.
    """
    eff = payload.get("effort") or {}
    mtg = payload.get("meetings") or {}

    def _money(v: Any) -> str:
        return f"${float(v):,.0f}" if v else "not recorded"

    rows = [["Field", "Value"]]
    rows.append(["Project", f"{payload.get('name') or payload.get('code')}"])
    rows.append(["Code", str(payload.get("code") or "")])
    rows.append(["Status", str(payload.get("status") or "unknown")])
    rows.append(["Account manager", str(payload.get("account_manager") or "—")])

    weeks = payload.get("duration_weeks")
    start = payload.get("start_date") or "?"
    rows.append([
        "Duration",
        f"{weeks} weeks (from {start})" if weeks else "not recorded",
    ])
    rows.append(["Budget", _money(payload.get("budget"))])
    tgt = payload.get("target_profit_pct")
    rows.append(["Target margin", f"{float(tgt):.0f}%" if tgt else "not recorded"])

    if eff.get("verified") and eff.get("total_hours"):
        rows.append([
            "Hours (allocated)",
            f"{eff['total_hours']} across {len(eff.get('by_person') or [])} "
            f"people, {eff.get('weeks', 0)} weeks",
        ])
        bph = payload.get("budget_per_hour")
        if bph:
            rows.append(["Budget ÷ hours", f"${bph}"])
    else:
        # Loud, because a silent blank here is what produced a retro that
        # said "hours not captured" while the table held 326.5 of them.
        rows.append(["Hours", "NOT READ — do not state a margin"])

    if mtg.get("count"):
        rows.append([
            "Meetings",
            f"{mtg['count']} · {mtg.get('total_hours', 0)}h total",
        ])
        share = mtg.get("tail_share")
        if share is not None:
            rows.append([
                f"Meeting time in final {mtg.get('tail_days', 14)} days",
                f"{float(share):.0%}",
            ])
    return rows


def effort_table(payload: dict) -> list[list[str]]:
    """Per-person allocated hours. Empty list when unread — the caller skips
    the section rather than rendering a table of zeros."""
    eff = payload.get("effort") or {}
    people = eff.get("by_person") or []
    if not eff.get("verified") or not people:
        return []
    rows = [["Person", "Hours (allocated)"]]
    rows.extend([[p.get("name", "?"), str(p.get("hours", 0))] for p in people])
    return rows
